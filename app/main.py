import gradio as gr
import PyPDF2
import docx
import re
import os
import base64
import hashlib
import hmac
import json
from urllib.parse import urlparse
import ssl
from datetime import datetime
from time import mktime
from urllib.parse import urlencode
from wsgiref.handlers import format_date_time
import websocket
import tempfile
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
import threading

# 讯飞星火认知大模型配置
SPARK_URL = 'wss://spark-api.xf-yun.com/v3.5/chat'
SPARK_APPID = os.getenv('SPARK_APPID', '35943c2b')
SPARK_API_SECRET = os.getenv('SPARK_API_SECRET', 'MmY4ZTBiYzZhNTJhOTMzMTY5MWZkNmFi')
SPARK_API_KEY = os.getenv('SPARK_API_KEY', 'd8c7001937dc33d1ec74aef5030ed816')


class SparkAPI:
    def __init__(self, appid, api_key, api_secret, spark_url):
        self.appid = appid
        self.api_key = api_key
        self.api_secret = api_secret
        self.spark_url = spark_url
        self.answer = ""

    def create_url(self):
        now = datetime.now()
        date = format_date_time(mktime(now.timetuple()))
        host = urlparse(self.spark_url).netloc
        path = urlparse(self.spark_url).path

        signature_origin = f"host: {host}\ndate: {date}\nGET {path} HTTP/1.1"
        signature_sha = hmac.new(self.api_secret.encode('utf-8'), signature_origin.encode('utf-8'),
                                 digestmod=hashlib.sha256).digest()
        signature_sha_base64 = base64.b64encode(signature_sha).decode(encoding='utf-8')
        authorization_origin = f'api_key="{self.api_key}", algorithm="hmac-sha256", headers="host date request-line", signature="{signature_sha_base64}"'
        authorization = base64.b64encode(authorization_origin.encode('utf-8')).decode(encoding='utf-8')
        v = {
            "authorization": authorization,
            "date": date,
            "host": host
        }
        return self.spark_url + '?' + urlencode(v)

    def on_message(self, ws, message):
        data = json.loads(message)
        code = data['header']['code']
        if code != 0:
            print(f'请求错误: {code}, {data}')
            ws.close()
        else:
            choices = data["payload"]["choices"]
            status = choices["status"]
            content = choices["text"][0]["content"]
            self.answer += content
            if status == 2:
                ws.close()

    def on_error(self, ws, error):
        print("### error:", error)

    def on_close(self, ws, close_status_code, close_msg):
        print("### closed ###")

    def on_open(self, ws):
        def run(*args):
            data = {
                "header": {
                    "app_id": self.appid
                },
                "parameter": {
                    "chat": {
                        "domain": "generalv3.5",
                        "temperature": 0.5,
                        "max_tokens": 1024,
                    }
                },
                "payload": {
                    "message": {
                        "text": [
                            {"role": "user", "content": self.question}
                        ]
                    }
                }
            }
            ws.send(json.dumps(data))

        threading.Thread(target=run).start()

    def call_api(self, question):
        self.question = question
        self.answer = ""
        websocket.enableTrace(False)
        ws_url = self.create_url()
        ws = websocket.WebSocketApp(ws_url,
                                    on_message=self.on_message,
                                    on_error=self.on_error,
                                    on_close=self.on_close,
                                    on_open=self.on_open)
        ws.run_forever(sslopt={"cert_reqs": ssl.CERT_NONE})
        return self.answer


# 初始化讯飞星火API
spark_api = SparkAPI(SPARK_APPID, SPARK_API_KEY, SPARK_API_SECRET, SPARK_URL)


def extract_content(file):
    content = ""
    if file.name.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(file.name)
        for page in pdf_reader.pages:
            content += page.extract_text()
    elif file.name.endswith('.docx'):
        doc = docx.Document(file.name)
        for para in doc.paragraphs:
            content += para.text + "\n"
    else:
        return "不支持的文件格式。请上传PDF或Word文档。"
    return content


def extract_fields(content):
    fields = {
        "姓名": re.search(r'姓名[：:]\s*(\S+)', content),
        "性别": re.search(r'性别[：:]\s*(\S+)', content),
        "出生年月": re.search(r'出生年月[：:]\s*(\S+)', content),
        "电话": re.search(r'电话[：:]\s*(\d{3}[-.]?\d{4}[-.]?\d{4})', content),
        "邮件": re.search(r'邮件[：:]\s*(\S+@\S+)', content),
        "期望工作地": re.search(r'期望工作地[：:]\s*(\S+)', content),
        "政治面貌": re.search(r'政治面貌[：:]\s*(\S+)', content),
        "籍贯": re.search(r'籍贯[：:]\s*(\S+)', content)
    }
    return {k: v.group(1) if v else "" for k, v in fields.items()}


def process_document(file):
    if file is None:
        return "请上传文件", "", "", "", "", "", "", "", ""
    content = extract_content(file)
    fields = extract_fields(content)
    return (
        content,
        fields["姓名"],
        fields["性别"],
        fields["出生年月"],
        fields["电话"],
        fields["邮件"],
        fields["期望工作地"],
        fields["政治面貌"],
        fields["籍贯"]
    )


def update_fields(content):
    fields = extract_fields(content)
    return (
        fields["姓名"],
        fields["性别"],
        fields["出生年月"],
        fields["电话"],
        fields["邮件"],
        fields["期望工作地"],
        fields["政治面貌"],
        fields["籍贯"]
    )


def add_education(school, degree, major, time, experience):
    return f"{school}, {degree}, {major}, {time}, {experience}"


def add_experience(time, company, position, content):
    return f"{time}, {company}, {position}, {content}"


def add_skill(time, name, description):
    return f"{time}, {name}, {description}"


def ai_optimize(text, field):
    prompt = f"请优化以下{field}，使其更加专业和有吸引力：\n\n{text}"
    return spark_api.call_api(prompt)


def optimize_in_school_exp(in_school_exp):
    return ai_optimize(in_school_exp, "在校经历")


def optimize_job_content(job_content):
    return ai_optimize(job_content, "工作内容")


def optimize_skill_description(skill_description):
    return ai_optimize(skill_description, "技能描述")


def optimize_self_evaluation(self_evaluation):
    return ai_optimize(self_evaluation, "自我评价")


def generate_resume(name, gender, birth, phone, email, location, status, hometown, education, experience, skills,
                    self_eval):
    resume = f"""
个人信息:
姓名：{name}
性别：{gender}
出生年月：{birth}
电话：{phone}
邮箱：{email}
期望工作地：{location}
政治面貌：{status}
籍贯：{hometown}

教育经历:
{education}

实践经验:
{experience}

技能和获奖情况:
{skills}

自我评价:
{self_eval}
    """
    return resume


def export_resume_pdf(name, gender, birth, phone, email, location, status, hometown,
                      education, experience, skills, self_eval):
    # 创建一个临时文件来保存PDF
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
        doc = SimpleDocTemplate(temp_file.name, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # 添加标题
        title_style = ParagraphStyle(name='Title', parent=styles['Heading1'], alignment=TA_CENTER)
        story.append(Paragraph(f"{name}的简历", title_style))
        story.append(Spacer(1, 12))

        # 添加个人信息
        story.append(Paragraph("个人信息", styles['Heading2']))
        info = f"性别：{gender}<br/>出生年月：{birth}<br/>电话：{phone}<br/>邮箱：{email}<br/>期望工作地：{location}<br/>政治面貌：{status}<br/>籍贯：{hometown}"
        story.append(Paragraph(info, styles['Normal']))
        story.append(Spacer(1, 12))

        # 添加教育经历
        story.append(Paragraph("教育经历", styles['Heading2']))
        story.append(Paragraph(education, styles['Normal']))
        story.append(Spacer(1, 12))

        # 添加实践经验
        story.append(Paragraph("实践经验", styles['Heading2']))
        story.append(Paragraph(experience, styles['Normal']))
        story.append(Spacer(1, 12))

        # 添加技能和获奖情况
        story.append(Paragraph("技能和获奖情况", styles['Heading2']))
        story.append(Paragraph(skills, styles['Normal']))
        story.append(Spacer(1, 12))

        # 添加自我评价
        story.append(Paragraph("自我评价", styles['Heading2']))
        story.append(Paragraph(self_eval, styles['Normal']))

        doc.build(story)

    return temp_file.name


def export_resume_word(name, gender, birth, phone, email, location, status, hometown,
                       education, experience, skills, self_eval):
    # 创建一个临时文件来保存Word文档
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        doc = docx.Document()

        # 添加标题
        doc.add_heading(f"{name}的简历", 0)

        # 添加个人信息
        doc.add_heading("个人信息", level=1)
        doc.add_paragraph(f"性别：{gender}")
        doc.add_paragraph(f"出生年月：{birth}")
        doc.add_paragraph(f"电话：{phone}")
        doc.add_paragraph(f"邮箱：{email}")
        doc.add_paragraph(f"期望工作地：{location}")
        doc.add_paragraph(f"政治面貌：{status}")
        doc.add_paragraph(f"籍贯：{hometown}")

        # 添加教育经历
        doc.add_heading("教育经历", level=1)
        doc.add_paragraph(education)

        # 添加实践经验
        doc.add_heading("实践经验", level=1)
        doc.add_paragraph(experience)

        # 添加技能和获奖情况
        doc.add_heading("技能和获奖情况", level=1)
        doc.add_paragraph(skills)

        # 添加自我评价
        doc.add_heading("自我评价", level=1)
        doc.add_paragraph(self_eval)

        doc.save(temp_file.name)

    return temp_file.name


with gr.Blocks() as iface:
    gr.Markdown("# AI职升姬v3")

    with gr.Row():
        with gr.Column():
            file_input = gr.File(label="上传简历文件（PDF或DOCX）")
            extract_button = gr.Button("提取信息")

        with gr.Column():
            content_output = gr.Textbox(label="提取的内容", lines=10, interactive=True)

    gr.Markdown("## 个人信息")
    with gr.Row():
        name = gr.Textbox(label="姓名")
        gender = gr.Radio(["男", "女"], label="性别")
        birth = gr.Textbox(label="出生年月")
        phone = gr.Textbox(label="电话")

    with gr.Row():
        email = gr.Textbox(label="邮箱")
        location = gr.Textbox(label="期望工作地")
        status = gr.Textbox(label="政治面貌")
        hometown = gr.Textbox(label="籍贯")

    gr.Markdown("## 教育经历")
    with gr.Row():
        school = gr.Textbox(label="学校名称")
        degree = gr.Dropdown(label="学历", choices=["大专", "本科", "研究生", "博士"])
        major = gr.Textbox(label="专业")
        edu_time = gr.Textbox(label="时间")

    with gr.Row():
        in_school_exp = gr.Textbox(label="在校经历", lines=3)
        optimize_in_school_exp_btn = gr.Button("AI优化在校经历")

    add_edu_btn = gr.Button("添加教育经历")
    edu_output = gr.Textbox(label="已添加的教育经历", lines=5)

# ... [前面的代码保持不变] ...

    gr.Markdown("## 实践经验")
    with gr.Row():
        exp_time = gr.Textbox(label="时间")
        company = gr.Textbox(label="公司/组织")
        position = gr.Textbox(label="职位")

    with gr.Row():
        job_content = gr.Textbox(label="工作内容", lines=3)
        optimize_job_content_btn = gr.Button("AI优化工作内容")

    add_exp_btn = gr.Button("添加实践经验")
    exp_output = gr.Textbox(label="已添加的实践经验", lines=5)

    gr.Markdown("## 技能和获奖情况")
    with gr.Row():
        skill_time = gr.Textbox(label="时间")
        skill_name = gr.Textbox(label="技能/奖项名称")

    with gr.Row():
        skill_description = gr.Textbox(label="描述", lines=3)
        optimize_skill_description_btn = gr.Button("AI优化技能描述")

    add_skill_btn = gr.Button("添加技能/奖项")
    skill_output = gr.Textbox(label="已添加的技能和获奖情况", lines=5)

    gr.Markdown("## 自我评价")
    with gr.Row():
        self_evaluation = gr.Textbox(label="自我评价", lines=5)
        optimize_self_evaluation_btn = gr.Button("AI优化自我评价")

    generate_btn = gr.Button("生成简历")
    resume_output = gr.Textbox(label="生成的简历", lines=20)

    export_pdf_btn = gr.Button("导出PDF")
    pdf_output = gr.File(label="导出的PDF简历")

    export_word_btn = gr.Button("导出Word")
    word_output = gr.File(label="导出的Word简历")

    extract_button.click(
        process_document,
        inputs=[file_input],
        outputs=[
            content_output,
            name,
            gender,
            birth,
            phone,
            email,
            location,
            status,
            hometown
        ]
    )

    content_output.change(
        update_fields,
        inputs=[content_output],
        outputs=[
            name,
            gender,
            birth,
            phone,
            email,
            location,
            status,
            hometown
        ]
    )

    add_edu_btn.click(
        add_education,
        inputs=[school, degree, major, edu_time, in_school_exp],
        outputs=edu_output
    )

    add_exp_btn.click(
        add_experience,
        inputs=[exp_time, company, position, job_content],
        outputs=exp_output
    )

    add_skill_btn.click(
        add_skill,
        inputs=[skill_time, skill_name, skill_description],
        outputs=skill_output
    )

    generate_btn.click(
        generate_resume,
        inputs=[name, gender, birth, phone, email, location, status, hometown, edu_output, exp_output, skill_output,
                self_evaluation],
        outputs=resume_output
    )

    optimize_in_school_exp_btn.click(
        optimize_in_school_exp,
        inputs=[in_school_exp],
        outputs=[in_school_exp]
    )

    optimize_job_content_btn.click(
        optimize_job_content,
        inputs=[job_content],
        outputs=[job_content]
    )

    optimize_skill_description_btn.click(
        optimize_skill_description,
        inputs=[skill_description],
        outputs=[skill_description]
    )

    optimize_self_evaluation_btn.click(
        optimize_self_evaluation,
        inputs=[self_evaluation],
        outputs=[self_evaluation]
    )

    export_pdf_btn.click(
        export_resume_pdf,
        inputs=[name, gender, birth, phone, email, location, status, hometown, edu_output, exp_output, skill_output,
                self_evaluation],
        outputs=pdf_output
    )

    export_word_btn.click(
        export_resume_word,
        inputs=[name, gender, birth, phone, email, location, status, hometown, edu_output, exp_output, skill_output,
                self_evaluation],
        outputs=word_output
    )

if __name__ == "__main__":
    iface.launch(server_name="0.0.0.0", share=False)